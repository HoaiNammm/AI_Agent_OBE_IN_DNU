"""
utils/dcct_parser.py

Trích xuất thông tin từ DCCT cũ (Word/PDF/Text/Markdown)
để tái chuẩn hóa theo chuẩn OBE/AUN-QA.
"""

import io
import json
import os
import re


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Đọc text thuần từ file upload.
    Hỗ trợ: .docx, .pdf, .txt, .md
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    if ext in ("txt", "md"):
        for enc in ("utf-8", "utf-16", "cp1258", "latin-1"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if ext == "docx":
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))
            return "\n".join(lines)
        except Exception:
            return ""

    if ext == "pdf":
        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
            return "\n".join(pages)
        except Exception:
            return ""

    return ""


def parse_dcct_info(raw_text: str) -> dict:
    """
    Trích xuất thông tin cốt lõi từ DCCT cũ.
    Ưu tiên Gemini Flash; fallback sang regex nếu không có API key.

    Returns:
        dict với keys: course_code, course_name, credits, summary, outline
    """
    api_key = os.environ.get("GOOGLE_API_KEY", "")
    if api_key:
        result = _parse_with_gemini(raw_text, api_key)
        if result.get("course_name") or result.get("course_code"):
            return result
    return _parse_with_regex(raw_text)


def _parse_with_gemini(raw_text: str, api_key: str) -> dict:
    try:
        import google.generativeai as genai

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")

        # Giới hạn token đầu vào
        truncated = raw_text[:10000]

        prompt = f"""Bạn là trợ lý phân tích đề cương học phần đại học Việt Nam.

Trích xuất các thông tin sau từ đề cương phía dưới:
1. course_code : Mã học phần (VD: CSC4007, MATH101, IT001...)
2. course_name : Tên học phần đầy đủ (tiếng Việt ưu tiên)
3. credits     : Số tín chỉ (chỉ số nguyên 1-5, không có đơn vị)
4. summary     : Đoạn mô tả mục tiêu/nội dung học phần (100-250 từ, trích nguyên văn từ tài liệu)
5. outline     : Danh sách các buổi/tuần học kèm chủ đề theo format:
                 "Buổi 1: [chủ đề]\\nBuổi 2: [chủ đề]\\n..."
                 (Trích thực tế từ đề cương, KHÔNG tự bịa thêm nội dung)

Trả về JSON duy nhất, KHÔNG có markdown fence hay giải thích:
{{"course_code":"...","course_name":"...","credits":"3","summary":"...","outline":"..."}}

Nếu không tìm thấy thông tin, để chuỗi rỗng "".

=== NỘI DUNG ĐỀ CƯƠNG ===
{truncated}"""

        response = model.generate_content(prompt)
        return _safe_json_parse(response.text)
    except Exception:
        return {}


def _safe_json_parse(text: str) -> dict:
    """Parse JSON từ LLM response, chịu được markdown fences."""
    empty = {
        "course_code": "",
        "course_name": "",
        "credits": "3",
        "summary": "",
        "outline": "",
    }
    # Bỏ markdown fences nếu có
    clean = re.sub(r"```(?:json)?", "", text).strip().rstrip("`").strip()
    # Tìm block JSON đầu tiên
    match = re.search(r"\{.*\}", clean, re.DOTALL)
    if match:
        try:
            data = json.loads(match.group())
            # Chuẩn hoá credits thành string
            if isinstance(data.get("credits"), int):
                data["credits"] = str(data["credits"])
            return {**empty, **data}
        except Exception:
            pass
    return empty


def _parse_with_regex(raw_text: str) -> dict:
    """Fallback: trích xuất cơ bản bằng regex, không cần LLM."""
    result = {
        "course_code": "",
        "course_name": "",
        "credits": "3",
        "summary": "",
        "outline": "",
    }

    # Mã học phần
    code_match = re.search(r"\b([A-Z]{2,5}\s*\d{3,5})\b", raw_text)
    if code_match:
        result["course_code"] = code_match.group(1).replace(" ", "")

    # Số tín chỉ
    for pat in [
        r"(\d)\s*tín\s*chỉ",
        r"Số tín chỉ[:\s]+(\d)",
        r"Credit[s]?[:\s]+(\d)",
    ]:
        m = re.search(pat, raw_text, re.IGNORECASE)
        if m and m.group(1) in "2345":
            result["credits"] = m.group(1)
            break

    # Mô tả: tìm section có từ khoá mô tả/mục tiêu
    desc_match = re.search(
        r"(?:mô tả|mục tiêu|giới thiệu)[^\n]*\n([\s\S]{50,600}?)(?=\n[A-ZĐÂĂÊÔƠƯÁÀẢÃẠ\d]|\Z)",
        raw_text,
        re.IGNORECASE,
    )
    if desc_match:
        result["summary"] = desc_match.group(1).strip()[:500]
    else:
        for line in raw_text.splitlines():
            if len(line.strip()) > 60:
                result["summary"] = line.strip()[:400]
                break

    # Sườn nội dung: tìm dòng có "Buổi" hoặc "Tuần" hoặc "Chương"
    outline_lines = []
    for line in raw_text.splitlines():
        if re.match(r"^\s*(Buổi|Tuần|Chương|Session|Week|Chapter)\s*\d+", line, re.IGNORECASE):
            outline_lines.append(line.strip())
    if outline_lines:
        result["outline"] = "\n".join(outline_lines)

    return result
