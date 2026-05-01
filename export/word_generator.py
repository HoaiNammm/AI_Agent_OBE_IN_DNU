"""
Word Generator - Xuất file DCCT theo định dạng Word chuẩn của Khoa CNTT - ĐH Đại Nam
Sử dụng python-docx để tạo tài liệu Word theo đúng mẫu DCCT của trường.
"""

import os
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path

from utils.logger import get_logger

logger = get_logger("export.word_generator")

OUTPUT_DIR = Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Tên đầy đủ các chương trình đào tạo
PROGRAM_NAMES = {
    "KHMT": "Khoa học máy tính",
    "HTTT": "Hệ thống thông tin",
    "GENERIC": "Công nghệ thông tin",
}

# Map hình thức dạy (tiếng Anh) → Hình thức giảng dạy (tiếng Việt)
FORM_MAP = {
    "Lecture": "Giảng giải có cấu trúc",
    "Lab": "Thực hành có hướng dẫn",
    "Discussion": "Thảo luận",
    "Exam": "Kiểm tra",
    "Presentation": "Báo cáo/Thuyết trình",
    "Project": "Thực hiện dự án",
    "LT": "Giảng giải có cấu trúc",
    "TH": "Thực hành có hướng dẫn",
    "BT": "Bài tập",
}


async def export_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Xuất DCCT ra file Word theo mẫu chuẩn Đại Nam.

    Input state: final_dcct_data hoặc dữ liệu riêng lẻ
    Output state: export_ready=True, export_path
    """
    logger.info("[Export] Bắt đầu xuất file Word DCCT (format Đại Nam)")

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError as e:
        logger.error(f"[Export] python-docx chưa được cài đặt: {e}")
        return {
            "export_ready": False,
            "errors": state.get("errors", []) + ["Export: python-docx chưa được cài đặt"],
        }

    # Lấy dữ liệu DCCT
    dcct_data = state.get("final_dcct_data") or _build_dcct_from_state(state)
    course_info = dcct_data.get("course_info", {})

    try:
        doc = Document()
        _set_document_style(doc)

        # ====== HEADER + TIÊU ĐỀ ======
        _add_dnu_header(doc, course_info)

        # ====== I. THÔNG TIN TỔNG QUÁT ======
        _add_section_i_general_info(doc, course_info, state)

        # ====== II. THÔNG TIN GIẢNG VIÊN ======
        _add_section_ii_instructor(doc, dcct_data.get("instructor_info", {}))

        # ====== III. TÓM TẮT NỘI DUNG ======
        _add_section_iii_summary(doc, state)

        # ====== IV. MỤC TIÊU HỌC PHẦN ======
        _add_section_iv_objectives(doc, dcct_data.get("objectives", []), dcct_data.get("clo_list", []))

        # ====== V. CHUẨN ĐẦU RA HỌC PHẦN (CLO) ======
        _add_section_v_clo(doc, dcct_data.get("clo_list", []))

        # ====== VI. TÀI LIỆU HỌC TẬP ======
        _add_section_vi_references(doc, dcct_data.get("references", []))

        # ====== VII. KẾ HOẠCH GIẢNG DẠY ======
        _add_section_vii_teaching_plan(doc, dcct_data.get("teaching_plan", []))

        # ====== VIII. ĐÁNH GIÁ HỌC PHẦN ======
        _add_section_viii_assessment(
            doc,
            dcct_data.get("assessment_plan", []),
            dcct_data.get("rubrics", {}),
            dcct_data.get("clo_list", []),
        )

        # ====== IX. ĐIỀU KIỆN THỰC HIỆN ======
        _add_section_ix_conditions(doc)

        # ====== KÝ TÊN ======
        _add_signature_block(doc)

        # ====== LƯU FILE ======
        course_code = course_info.get("code", "DCCT")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"DCCT_{course_code}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename

        doc.save(str(filepath))
        logger.info(f"[Export] File xuất thành công: {filepath}")

        # ── Auto-index ĐCCT vào knowledge base Q&A ──────────────────────────
        try:
            from agents.qa_agent import index_dcct_from_state
            course_code_key = course_info.get("code", state.get("course_code", "UNKNOWN"))
            idx_info = index_dcct_from_state(course_code_key, state)
            logger.info(
                f"[Export] ĐCCT indexed vào Q&A store: "
                f"{idx_info['chunks_indexed']} chunks cho {course_code_key}"
            )
        except Exception as idx_err:
            logger.warning(f"[Export] Không thể index ĐCCT vào Q&A store: {idx_err}")

        return {
            "export_ready": True,
            "export_path": str(filepath),
        }

    except Exception as e:
        logger.error(f"[Export] Lỗi xuất Word: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            "export_ready": False,
            "errors": state.get("errors", []) + [f"Export: {e}"],
        }


# ============================================================
# DOCUMENT STYLE
# ============================================================

def _set_document_style(doc):
    """Thiết lập style mặc định cho tài liệu."""
    from docx.shared import Pt, Cm

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)


# ============================================================
# SECTION I-IX: NỘI DUNG CHÍNH
# ============================================================

def _add_dnu_header(doc, course_info: Dict):
    """Thêm header 2 cột + tiêu đề theo mẫu ĐH Đại Nam."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ── Bảng header 2 cột ──────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    # Ẩn border bảng header
    _remove_table_borders(tbl)

    left_cell = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)

    # Cột trái: tên trường + khoa
    _add_cell_paragraph(left_cell, "TRƯỜNG ĐẠI HỌC ĐẠI NAM", bold=True, size=12,
                        align=WD_ALIGN_PARAGRAPH.CENTER, clear_existing=True)
    _add_cell_paragraph(left_cell, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, size=12,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    # Cột phải: quốc hiệu
    _add_cell_paragraph(right_cell, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=12,
                        align=WD_ALIGN_PARAGRAPH.CENTER, clear_existing=True)
    _add_cell_paragraph(right_cell, "Độc lập - Tự do - Hạnh phúc", bold=True, size=12,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_cell_paragraph(right_cell, "───────────────────", bold=False, size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Tiêu đề chính ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True

    # ── Tên học phần ──────────────────────────────────────────────────────
    course_name = course_info.get("name", "")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Tên học phần: {course_name}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(13)
    r2.font.bold = True

    # ── Trình độ đào tạo ───────────────────────────────────────────────────
    cohort = course_info.get("cohort", "20")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p3.add_run(f"Trình độ đào tạo: Đại học chính quy, áp dụng từ khóa: {cohort}")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)
    r3.font.bold = True

    doc.add_paragraph()


def _add_section_i_general_info(doc, course_info: Dict, state: Dict):
    """I. Thông tin tổng quát."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _add_bold_heading(doc, "I. Thông tin tổng quát")

    credits_str = str(course_info.get("credits", "3"))
    try:
        credits_int = int(credits_str)
    except ValueError:
        credits_int = 3

    theory_periods = int(course_info.get("theory_periods", 0) or 0)
    lab_periods = int(course_info.get("lab_periods", 0) or 0)
    if theory_periods == 0 and lab_periods == 0:
        # Tính mặc định: 1/3 LT, 2/3 TH cho 3 TC
        theory_tc = max(1, credits_int // 3)
        lab_tc = credits_int - theory_tc
        theory_periods = theory_tc * 15
        lab_periods = lab_tc * 15
    self_study = credits_int * 30 + 15  # heuristic

    name_vn = course_info.get("name", "")
    name_en = course_info.get("name_en", course_info.get("name_english", ""))
    code = course_info.get("code", "")
    program_code = state.get("program") or course_info.get("program", "")
    program = PROGRAM_NAMES.get(program_code, program_code or "Khoa học máy tính")
    prerequisites = course_info.get("prerequisites", [])
    prereq_str = ", ".join(prerequisites) if prerequisites else "Không có"
    department = course_info.get("department", "Khoa Công nghệ thông tin")
    issue_date = course_info.get("issue_date", datetime.now().strftime("%d/%m/%Y"))
    issue_number = course_info.get("issue_number", "1")

    theory_tc = theory_periods // 15
    lab_tc = lab_periods // 15

    items = [
        ("1.", f"Tên học phần (tiếng Việt): {name_vn}"),
        ("2.", f"Tên học phần (tiếng Anh): {name_en}"),
        ("3.", f"Mã học phần: {code}"),
        ("4.", f"Thuộc chương trình đào tạo ngành: {program}"),
        ("5.", f"Số tín chỉ: {credits_int} ({theory_periods}, {lab_periods})"),
        ("6.", "Hoạt động học tập"),
        ("", f"- Lý thuyết, bài tập (LT, BT): {theory_tc:02d} TC ({theory_periods} giờ)"),
        ("", f"- Thực hành, thí nghiệm (TH): {lab_tc:02d} TC ({lab_periods} giờ)"),
        ("", f"- Tự học: {self_study} giờ"),
        ("7.", f"Học phần học trước: {prereq_str}"),
        ("8.", f"Bộ môn (Khoa phụ trách): {department}"),
        ("9.", f"Ngày ban hành: {issue_date}"),
        ("10.", f"Lần ban hành: {issue_number}"),
    ]

    for num, text in items:
        p = doc.add_paragraph()
        if num:
            run_num = p.add_run(f"{num} ")
            run_num.font.name = "Times New Roman"
            run_num.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(12)
            run_num.bold = False
        run_txt = p.add_run(text)
        run_txt.font.name = "Times New Roman"
        run_txt.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(12)
        if not num:
            p.paragraph_format.left_indent = __import__("docx.shared", fromlist=["Cm"]).Cm(1)

    doc.add_paragraph()


def _add_section_ii_instructor(doc, instructor_info: Dict):
    """II. Thông tin giảng viên."""
    _add_bold_heading(doc, "II. Thông tin giảng viên")

    headers = ["TT", "Họ và tên", "Chức danh, học vị", "Liên hệ", "Đơn vị công tác", "Ghi chú"]
    table = doc.add_table(rows=3, cols=len(headers))
    table.style = "Table Grid"

    # Hàng header
    for i, h in enumerate(headers):
        _set_cell_text(table.cell(0, i), h, bold=True, center=True)

    # Hàng phụ: nhóm "Giảng viên phụ trách học phần"
    # Merge toàn bộ hàng 1
    merged = table.cell(1, 0).merge(table.cell(1, len(headers) - 1))
    _set_cell_text(merged, "I    Giảng viên phụ trách học phần", bold=True)

    # Hàng dữ liệu 1
    gv = instructor_info if instructor_info else {}
    _set_cell_text(table.cell(2, 0), "1")
    _set_cell_text(table.cell(2, 1), gv.get("name", ""))
    _set_cell_text(table.cell(2, 2), gv.get("title", ""))
    _set_cell_text(table.cell(2, 3), gv.get("contact", ""))
    _set_cell_text(table.cell(2, 4), gv.get("department", "Khoa CNTT"))
    _set_cell_text(table.cell(2, 5), gv.get("note", ""))

    doc.add_paragraph()


def _add_section_iii_summary(doc, state: Dict):
    """III. Tóm tắt nội dung học phần."""
    _add_bold_heading(doc, "III. Tóm tắt nội dung học phần")

    summary = state.get("summary", "")
    p = doc.add_paragraph(summary)
    _set_para_font(p, 12)
    doc.add_paragraph()


def _add_section_iv_objectives(doc, objectives: List, clo_list: List[Dict]):
    """IV. Mục tiêu của học phần."""
    _add_bold_heading(doc, "IV. Mục tiêu của học phần")

    # Nếu không có objectives riêng, sinh từ CLO
    if not objectives:
        objectives = [clo.get("description", "") for clo in clo_list]

    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {obj}")
        run.font.name = "Times New Roman"
        run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(12)

    doc.add_paragraph()


def _add_section_v_clo(doc, clo_list: List[Dict]):
    """V. Chuẩn đầu ra của học phần."""
    _add_bold_heading(doc, "V. Chuẩn đầu ra của học phần")

    if not clo_list:
        doc.add_paragraph("Chưa có CLO.")
        return

    headers = ["Mã CĐR\ncủa HP", "Nội dung CĐR của HP", "Mã tiêu chí đánh giá", "Mức độ"]
    table = doc.add_table(rows=1 + len(clo_list), cols=len(headers))
    table.style = "Table Grid"

    # Thiết lập độ rộng cột Section V
    from docx.shared import Cm
    v_col_widths = [Cm(2.0), Cm(8.5), Cm(3.5), Cm(2.0)]
    for ci, width in enumerate(v_col_widths):
        for cell in table.columns[ci].cells:
            cell.width = width

    for i, h in enumerate(headers):
        _set_cell_text(table.cell(0, i), h, bold=True, center=True)

    for row_idx, clo in enumerate(clo_list, 1):
        code = clo.get("code", f"CLO {row_idx}")
        description = clo.get("description", "")
        pi_codes = clo.get("pi_codes", [])
        pi_str = "; ".join(pi_codes) if pi_codes else ""
        irma = clo.get("mapping_level", "R")

        _set_cell_text(table.cell(row_idx, 0), code)
        _set_cell_text(table.cell(row_idx, 1), description)
        _set_cell_text(table.cell(row_idx, 2), pi_str)
        _set_cell_text(table.cell(row_idx, 3), irma, center=True)

    doc.add_paragraph()


def _add_section_vi_references(doc, references: List[Dict]):
    """VI. Tài liệu học tập."""
    _add_bold_heading(doc, "VI. Tài liệu học tập")

    # Bảng có header + merge ô "Mục đích sử dụng"
    # Cột: TT | Tên tài liệu | Tên tác giả | Năm XB | Nhà XB | Bắt buộc | Tham khảo
    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"

    # Row 0: header chính - merge cột 5+6 thành "Mục đích sử dụng"
    _set_cell_text(table.cell(0, 0), "TT", bold=True, center=True)
    _set_cell_text(table.cell(0, 1), "Tên tài liệu", bold=True, center=True)
    _set_cell_text(table.cell(0, 2), "Tên tác giả/Đơn vị", bold=True, center=True)
    _set_cell_text(table.cell(0, 3), "Năm XB", bold=True, center=True)
    _set_cell_text(table.cell(0, 4), "Nhà XB", bold=True, center=True)
    merged_purpose = table.cell(0, 5).merge(table.cell(0, 6))
    _set_cell_text(merged_purpose, "Mục đích sử dụng", bold=True, center=True)

    # Row 1: sub-header
    _set_cell_text(table.cell(1, 0), "", bold=True, center=True)
    _set_cell_text(table.cell(1, 1), "", bold=True, center=True)
    _set_cell_text(table.cell(1, 2), "", bold=True, center=True)
    _set_cell_text(table.cell(1, 3), "", bold=True, center=True)
    _set_cell_text(table.cell(1, 4), "", bold=True, center=True)
    _set_cell_text(table.cell(1, 5), "Tài liệu bắt buộc", bold=True, center=True)
    _set_cell_text(table.cell(1, 6), "Tài liệu tham khảo", bold=True, center=True)

    if not references:
        # Thêm 1 hàng rỗng nếu chưa có tài liệu
        row = table.add_row()
        _set_cell_text(row.cells[0], "")
        return

    for i, ref in enumerate(references, 1):
        row = table.add_row()
        cells = row.cells
        _set_cell_text(cells[0], str(i), center=True)
        _set_cell_text(cells[1], ref.get("title", ""))
        _set_cell_text(cells[2], ref.get("author", ""))
        _set_cell_text(cells[3], str(ref.get("year", "")), center=True)
        _set_cell_text(cells[4], ref.get("publisher", ""))
        is_required = ref.get("required", True)
        _set_cell_text(cells[5], "x" if is_required else "", center=True)
        _set_cell_text(cells[6], "x" if not is_required else "", center=True)

    doc.add_paragraph()


def _add_section_vii_teaching_plan(doc, teaching_plan: List[Dict]):
    """VII. Kế hoạch giảng dạy — 8 cột theo mẫu DNU."""
    _add_bold_heading(doc, "VII. Kế hoạch giảng dạy")

    if not teaching_plan:
        doc.add_paragraph("Chưa có kế hoạch giảng dạy.")
        return

    headers = [
        "Buổi học",
        "Số tiết",
        "Hình thức giảng dạy",
        "Nội dung",
        "Chuẩn đầu ra học phần",
        "Hoạt động dạy và học",
        "Bài\nđánh giá",
        "Yêu cầu người học chuẩn bị",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Thiết lập độ rộng cột (tổng ~16.5cm)
    from docx.shared import Cm
    col_widths = [Cm(1.2), Cm(1.3), Cm(2.8), Cm(4.5), Cm(1.8), Cm(2.5), Cm(1.0), Cm(2.5)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    for i, h in enumerate(headers):
        _set_cell_text(table.cell(0, i), h, bold=True, center=True)

    # Nhóm sessions theo số buổi (no) để gộp các phần LT+TH cùng buổi
    grouped: Dict[int, List[Dict]] = {}
    for session in teaching_plan:
        no = int(session.get("no", 0))
        grouped.setdefault(no, []).append(session)

    for no in sorted(grouped.keys()):
        sessions = grouped[no]
        for s_idx, session in enumerate(sessions):
            row = table.add_row()
            cells = row.cells

            is_main_row = (s_idx == 0)

            # Cột 0: Buổi học — chỉ hiển thị ở sub-row đầu tiên
            _set_cell_text(cells[0], str(no) if is_main_row else "", center=True)

            # Cột 1: Số tiết
            periods = session.get("duration_periods", session.get("periods", 1))
            _set_cell_text(cells[1], str(periods), center=True)

            # Cột 2: Hình thức giảng dạy
            # Ưu tiên dùng activities nếu là văn bản tiếng Việt mô tả đầy đủ,
            # nếu là từ tiếng Anh đơn giản thì map sang tiếng Việt
            activities = session.get("activities", "")
            stype = session.get("type", "LT")
            if activities and activities not in FORM_MAP:
                # Đã là mô tả tiếng Việt đầy đủ
                form_text = activities
            else:
                form_text = FORM_MAP.get(activities) or FORM_MAP.get(stype, stype)
            _set_cell_text(cells[2], form_text)

            # Cột 3: Nội dung
            content = session.get("content", "")
            details = session.get("details", "")
            content_full = f"{content}\n{details}".strip() if details else content
            _set_cell_text(cells[3], content_full)

            # Cột 4-7: chỉ điền ở main row (LT); TH sub-rows để trống
            if is_main_row:
                # Cột 4: Chuẩn đầu ra HP
                clo_codes = session.get("clo_codes", [])
                _set_cell_text(cells[4], "; ".join(clo_codes))

                # Cột 5: Hoạt động dạy và học
                _set_cell_text(cells[5], session.get("teaching_activities", ""))

                # Cột 6: Bài đánh giá
                _set_cell_text(cells[6], session.get("assessment", ""), center=True)

                # Cột 7: Yêu cầu chuẩn bị
                _set_cell_text(cells[7], session.get("preparation", ""))
            else:
                _set_cell_text(cells[4], "")
                _set_cell_text(cells[5], "")
                _set_cell_text(cells[6], "")
                _set_cell_text(cells[7], "")

    doc.add_paragraph()


def _add_section_viii_assessment(
    doc, assessment_plan: List[Dict], rubrics: Dict, clo_list: List[Dict]
):
    """VIII. Đánh giá học phần."""
    _add_bold_heading(doc, "VIII. Đánh giá học phần")

    # ── 8.1 Nội dung đánh giá ───────────────────────────────────────────────
    _add_sub_heading(doc, "8.1. Nội dung đánh giá")

    if not assessment_plan:
        doc.add_paragraph("Chưa có hệ thống đánh giá.")
    else:
        headers = [
            "Bài đánh giá",
            "Hình thức đánh giá",
            "Nội dung",
            "Thời điểm",
            "Tiêu chí đánh giá",
            "Trọng số %",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            _set_cell_text(table.cell(0, i), h, bold=True, center=True)

        total_weight = 0
        for a in assessment_plan:
            row = table.add_row()
            cells = row.cells
            weight = float(a.get("weight", 0))
            total_weight += weight
            _set_cell_text(cells[0], a.get("code", "") + ". " + a.get("name", ""))
            _set_cell_text(cells[1], a.get("format", ""))
            _set_cell_text(cells[2], a.get("description", ""))
            _set_cell_text(cells[3], a.get("timing", a.get("frequency", "")))
            _set_cell_text(cells[4], a.get("criteria_summary", ""))
            _set_cell_text(cells[5], f"{weight * 100:.0f}%", center=True)

        # Hàng tổng
        row = table.add_row()
        merged_total = row.cells[0].merge(row.cells[4])
        _set_cell_text(merged_total, "Tổng cộng", bold=True, center=True)
        _set_cell_text(row.cells[5], f"{total_weight * 100:.0f}%", bold=True, center=True)

        # Ghi chú
        p = doc.add_paragraph(
            "Ghi chú sử dụng trong hồ sơ: A1 là thành phần điểm quá trình; "
            "hiện không có căn cứ để dùng A1 làm học phần A hoặc làm minh chứng kết luận đạt CLO/PLO."
        )
        _set_para_font(p, 11, italic=True)

    doc.add_paragraph()

    # ── 8.2 Rubric đánh giá ─────────────────────────────────────────────────
    _add_sub_heading(doc, "8.2. Rubric đánh giá")

    # 8.2.1 Bảng truy vết
    _add_sub2_heading(doc, "8.2.1. Bảng truy vết assessment – CLO – PI/PLO")

    _add_truy_vet_table(doc, assessment_plan, clo_list)
    doc.add_paragraph()

    # 8.2.2 Bảng ma trận CLO × PLO
    _add_sub2_heading(doc, "8.2.2. Bảng ma trận CLO × PLO")
    _add_clo_plo_matrix_table(doc, clo_list)
    doc.add_paragraph()

    # 8.2.3+ Rubric từng cấu phần
    rubric_order = ["A1", "A2.1", "A2.2", "A3"]
    existing_codes = list(rubrics.keys())
    # Sắp theo thứ tự chuẩn, còn lại append sau
    ordered = [c for c in rubric_order if c in rubrics] + \
              [c for c in existing_codes if c not in rubric_order]

    for idx, component_code in enumerate(ordered):
        sub_idx = idx + 3  # 8.2.3, 8.2.4, ... (vì 8.2.2 là CLO×PLO matrix)
        sub_title = f"8.2.{sub_idx}. Rubric {component_code}"
        _add_sub2_heading(doc, sub_title)

        rubric = rubrics.get(component_code, {})
        criteria = rubric.get("criteria", [])

        if not criteria:
            doc.add_paragraph("Chưa có rubric.")
            continue

        headers = [
            "Tiêu chí / Criterion",
            "CLO đo lường",
            "M1\n0–<5",
            "M2\n5–<7",
            "M3\n7–<8",
            "M4\n8–<9",
            "M5\n9–10",
            "Trọng số",
        ]
        table = doc.add_table(rows=1 + len(criteria), cols=len(headers))
        table.style = "Table Grid"

        # Độ rộng cột rubric (tổng ~16.5cm)
        from docx.shared import Cm
        rubric_col_widths = [Cm(3.0), Cm(1.5), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(1.5)]
        for ci, width in enumerate(rubric_col_widths):
            for cell in table.columns[ci].cells:
                cell.width = width

        for i, h in enumerate(headers):
            _set_cell_text(table.cell(0, i), h, bold=True, center=True)

        for row_idx, criterion in enumerate(criteria, 1):
            levels = criterion.get("levels", {})
            weight_pct = f"{float(criterion.get('weight_in_component', 0)) * 100:.0f}%"
            clo_meas = criterion.get("clo_measured", criterion.get("clo_measure", ""))

            _set_cell_text(table.cell(row_idx, 0), criterion.get("criterion", ""))
            _set_cell_text(table.cell(row_idx, 1), clo_meas)
            _set_cell_text(table.cell(row_idx, 2), _get_level_desc(levels, "M1", "fail"))
            _set_cell_text(table.cell(row_idx, 3), _get_level_desc(levels, "M2", "pass"))
            _set_cell_text(table.cell(row_idx, 4), _get_level_desc(levels, "M3", "pass"))
            _set_cell_text(table.cell(row_idx, 5), _get_level_desc(levels, "M4", "good"))
            _set_cell_text(table.cell(row_idx, 6), _get_level_desc(levels, "M5", "excellent"))
            _set_cell_text(table.cell(row_idx, 7), weight_pct, center=True)

        doc.add_paragraph()

    # Lưu ý OBE
    p = doc.add_paragraph(
        "Lưu ý OBE: với A2.1, A2.2 và A3, các criterion \"báo cáo/minh chứng\" là criterion "
        "hỗ trợ evidence, không nên dùng riêng để kết luận đạt CLO."
    )
    _set_para_font(p, 11, italic=True)
    doc.add_paragraph()


def _add_truy_vet_table(doc, assessment_plan: List[Dict], clo_list: List[Dict]):
    """Bảng truy vết assessment – CLO – PI/PLO."""
    headers = [
        "Assessment",
        "Tỷ trọng",
        "CLO đo lường",
        "PI liên quan",
        "PLO liên quan",
        "Ghi chú truy vết",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.cell(0, i), h, bold=True, center=True)

    for a in assessment_plan:
        row = table.add_row()
        cells = row.cells
        weight = float(a.get("weight", 0))
        clo_raw = a.get("clo_mapping", [])
        # A1 không gắn CLO chính thức
        if a.get("code") == "A1" or not clo_raw:
            clo_list_str = "Không gắn CLO chính thức" if a.get("code") == "A1" else ""
        else:
            clo_list_str = ", ".join(clo_raw)
        pi_list_str = "; ".join(a.get("pi_codes", []))
        plo_list_str = "; ".join(a.get("plo_codes", []))
        truy_vet_note = a.get("truy_vet_note", a.get("note", ""))

        _set_cell_text(cells[0], a.get("code", ""))
        _set_cell_text(cells[1], f"{weight * 100:.0f}%", center=True)
        _set_cell_text(cells[2], clo_list_str)
        _set_cell_text(cells[3], pi_list_str)
        _set_cell_text(cells[4], plo_list_str)
        _set_cell_text(cells[5], truy_vet_note)


def _add_clo_plo_matrix_table(doc, clo_list: List[Dict]):
    """Bảng ma trận CLO × PLO với mức IRMA từ clo_list[*].plo_mapping."""
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Thu thập PLO duy nhất từ mapping
    plo_codes_set = set()
    for clo in clo_list:
        for pm in clo.get("plo_mapping", []):
            plo_codes_set.add(pm.get("plo_code", ""))
    if not plo_codes_set:
        # Thử từ pi_mapping → primary_plo
        for clo in clo_list:
            for pm in clo.get("pi_mapping", []):
                p = pm.get("primary_plo", "")
                if p:
                    plo_codes_set.add(p)
    plo_codes = sorted(plo_codes_set)
    if not plo_codes:
        return  # Không có dữ liệu mapping → bỏ qua

    # Header row: CLO | PLO1 | PLO2 | ...
    n_cols = 1 + len(plo_codes)
    table = doc.add_table(rows=1 + len(clo_list), cols=n_cols)
    table.style = "Table Grid"

    # Thiết lập độ rộng: cột CLO 3cm, mỗi PLO col = (tổng - 3) / n_plo
    clo_width = Cm(3.5)
    if len(plo_codes) > 0:
        plo_width = Cm(min(2.5, 12.0 / len(plo_codes)))
    else:
        plo_width = Cm(2.5)
    for cell in table.columns[0].cells:
        cell.width = clo_width
    for ci in range(1, n_cols):
        for cell in table.columns[ci].cells:
            cell.width = plo_width

    # Header
    _set_cell_text(table.cell(0, 0), "CLO \\ PLO", bold=True, center=True)
    for ci, plo in enumerate(plo_codes, start=1):
        _set_cell_text(table.cell(0, ci), plo, bold=True, center=True)

    # Xây dựng lookup: (clo_code, plo_code) → irma_level
    irma_lookup: Dict[tuple, str] = {}
    for clo in clo_list:
        code = clo.get("clo_code", "")
        for pm in clo.get("plo_mapping", []):
            irma_lookup[(code, pm.get("plo_code", ""))] = pm.get("irma_level", "x")
        for pm in clo.get("pi_mapping", []):
            plo = pm.get("primary_plo", "")
            irma = pm.get("irma_level", "x")
            if plo:
                key = (code, plo)
                if key not in irma_lookup:
                    irma_lookup[key] = irma

    # Điền dữ liệu
    for ri, clo in enumerate(clo_list, start=1):
        code = clo.get("clo_code", f"CLO{ri}")
        _set_cell_text(table.cell(ri, 0), code, bold=True)
        for ci, plo in enumerate(plo_codes, start=1):
            irma = irma_lookup.get((code, plo), "")
            _set_cell_text(table.cell(ri, ci), irma, center=True)


def _add_section_ix_conditions(doc):
    """IX. Điều kiện thực hiện học phần."""
    _add_bold_heading(doc, "IX. Điều kiện thực hiện học phần")

    _add_sub_heading(doc, "9.1. Quy định về tham dự lớp học")

    conditions_91 = [
        "Sinh viên có trách nhiệm tham dự đầy đủ các buổi học. Trong trường hợp nghỉ học có lý do "
        "bất khả kháng thì phải có minh chứng cụ thể và đầy đủ;",
        "Tham dự tối thiểu 80% số tiết trên lớp đối với nội dung/học phần lý thuyết; tham dự đủ 100% "
        "số tiết đối với nội dung/học phần thực hành tại Trường.",
    ]
    for c in conditions_91:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(c)
        run.font.name = "Times New Roman"
        run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(12)

    doc.add_paragraph()
    _add_sub_heading(doc, "9.2. Quy định về hành vi lớp học")

    conditions_92 = [
        "Học phần được thực hiện trên nguyên tắc tôn trọng người học và người dạy. "
        "Mọi hành vi làm ảnh hưởng đến quá trình dạy và học đều bị nghiêm cấm;",
        "Người học cần nghiên cứu giáo trình, chuẩn bị các ý kiến thảo luận; "
        "đọc và sưu tầm các tư liệu có liên quan đến nội dung của môn học;",
        "Người học tham dự đầy đủ, tích cực việc lên lớp, thảo luận, thi, kiểm tra theo quy định;",
        "Người học phải đi học đúng giờ theo quy định;",
        "Người học không được làm việc riêng, không được sử dụng điện thoại và đeo tai nghe trong giờ học;",
        "Người học không được ăn uống, không được nhai kẹo cao su, không được xả rác bừa bãi ra giảng đường;",
        "Quy định về nộp bài tập, bài kiểm tra: Người học nộp bài tập được giao đúng hạn và có mặt "
        "đúng giờ quy định để làm bài kiểm tra;",
        "Nội quy lớp học: Tuân theo nội quy của Trường Đại học Đại Nam.",
    ]
    for c in conditions_92:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(c)
        run.font.name = "Times New Roman"
        run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(12)

    doc.add_paragraph()


def _add_signature_block(doc):
    """Khối ký tên 3 cột cuối tài liệu."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=2, cols=3)
    _remove_table_borders(table)

    titles = [
        "TRUNG TÂM PHÁT TRIỂN CHƯƠNG TRÌNH ĐÀO TẠO",
        "TRƯỞNG KHOA/VIỆN TRƯỞNG",
        "ĐẠI DIỆN NHÓM BIÊN SOẠN",
    ]
    subtitles = ["(Ký, ghi rõ họ tên)"] * 3

    for i, (title, subtitle) in enumerate(zip(titles, subtitles)):
        _add_cell_paragraph(table.cell(0, i), title, bold=True, size=12,
                            align=WD_ALIGN_PARAGRAPH.CENTER, clear_existing=True)
        _add_cell_paragraph(table.cell(0, i), subtitle, bold=False, size=11,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        # Khoảng trống cho chữ ký
        _add_cell_paragraph(table.cell(1, i), "\n\n", bold=False, size=12,
                            align=WD_ALIGN_PARAGRAPH.CENTER, clear_existing=True)


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def _add_bold_heading(doc, text: str):
    """Thêm tiêu đề mục in đậm kiểu paragraph (không dùng Heading style)."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.bold = True


def _add_sub_heading(doc, text: str):
    """Thêm tiêu đề cấp 2 (8.x) in đậm nghiêng."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True


def _add_sub2_heading(doc, text: str):
    """Thêm tiêu đề cấp 3 (8.2.x) in đậm."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True


def _set_cell_text(cell, text: str, bold: bool = False, center: bool = False):
    """Xóa nội dung cell cũ và đặt text mới."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.bold = bold


def _add_cell_paragraph(cell, text: str, bold: bool = False, size: int = 12,
                        align=None, clear_existing: bool = False):
    """Thêm paragraph vào cell."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if clear_existing:
        for para in cell.paragraphs:
            para.clear()
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    else:
        para = cell.add_paragraph()

    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def _remove_table_borders(table):
    """Ẩn toàn bộ border của bảng."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _get_level_desc(levels: Dict, m_key: str, fallback_key: str) -> str:
    """Lấy mô tả mức rubric — hỗ trợ cả định dạng M1-M5 và excellent/good/pass/fail."""
    # Thử trực tiếp key M1..M5
    val = levels.get(m_key)
    if val:
        return val if isinstance(val, str) else val.get("description", "")

    # Fallback sang định dạng 4 mức cũ nếu LLM chưa cập nhật
    fallback = levels.get(fallback_key)
    if fallback:
        return fallback if isinstance(fallback, str) else fallback.get("description", "")
    return ""


def _set_para_font(para, size: int = 12, bold: bool = False, italic: bool = False):
    """Thiết lập font cho paragraph."""
    from docx.shared import Pt
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    if not para.runs:
        run = para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def _build_dcct_from_state(state: Dict) -> Dict:
    """Xây dựng DCCT data từ state nếu final_dcct_data chưa có."""
    extracted = state.get("extracted_info", {})
    return {
        "course_info": {
            "code": state.get("course_code", ""),
            "name": state.get("course_name", ""),
            "name_en": extracted.get("name_en", extracted.get("name_english", "")),
            "credits": state.get("credits", "3"),
            "theory_periods": extracted.get("theory_periods", 0),
            "lab_periods": extracted.get("lab_periods", 0),
            "prerequisites": extracted.get("prerequisites", []),
            "program": state.get("program", ""),
            "department": extracted.get("department", "Khoa Công nghệ thông tin"),
            "cohort": extracted.get("cohort", "20"),
            "issue_date": extracted.get("issue_date", ""),
            "issue_number": extracted.get("issue_number", "1"),
        },
        "instructor_info": state.get("instructor_info", {}),
        "objectives": state.get("objectives", []),
        "clo_list": state.get("clo_list", []),
        "mapping_matrix": state.get("mapping_matrix", []),
        "teaching_plan": state.get("teaching_plan", []),
        "assessment_plan": state.get("assessment_plan", []),
        "rubrics": state.get("rubrics", {}),
        "references": state.get("references", extracted.get("references", [])),
        "confidence_score": state.get("confidence_score", 0),
    }
