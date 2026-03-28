"""
word_generator.py - Sinh file DCCT Word từ template.

Generates a properly formatted DCCT (Đề Cương Chi Tiết) Word document
using the official university template (DCCT_Template.docx).
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import TEMPLATE_PATH, OUTPUT_DIR

if TYPE_CHECKING:
    from state import DCCTState

logger = logging.getLogger(__name__)


def _get_or_create_doc() -> Document:
    """Load template if available, otherwise start with a blank document."""
    template = Path(TEMPLATE_PATH)
    if template.exists():
        logger.info("[WordGenerator] Using template: %s", TEMPLATE_PATH)
        return Document(str(template))
    logger.warning("[WordGenerator] Template not found, using blank document.")
    return Document()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_table_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = str(cell_text)


def generate_word_document(state: "DCCTState") -> str:
    """
    Generate a DCCT Word document from the pipeline state.

    Args:
        state: Completed DCCTState with all fields populated.

    Returns:
        Absolute path to the generated .docx file.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = _get_or_create_doc()

    course_name = state.get("course_name", "Unknown")
    course_code = state.get("course_code", "N/A")
    credits = state.get("credits", "N/A")
    department = state.get("department", "N/A")

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Tên học phần: {course_name}")
    doc.add_paragraph(f"Mã học phần: {course_code}")
    doc.add_paragraph(f"Số tín chỉ: {credits}")
    doc.add_paragraph(f"Khoa: {department}")
    doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_paragraph("")

    # ── CLOs ──────────────────────────────────────────────────────────────────
    clos = state.get("clos", [])
    if clos:
        _add_heading(doc, "1. Chuẩn Đầu Ra Học Phần (CLOs)", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Mã CLO"
        hdr[1].text = "Mô tả"
        hdr[2].text = "Cấp độ Bloom"
        hdr[3].text = "Động từ hành động"
        for clo in clos:
            _add_table_row(
                tbl,
                [clo.id, clo.description, clo.bloom_level, ", ".join(clo.verbs)],
            )
        doc.add_paragraph("")

    # ── PLOs ──────────────────────────────────────────────────────────────────
    plos = state.get("plos", [])
    if plos:
        _add_heading(doc, "2. Ma Trận CLO–PLO", level=1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Mã PLO"
        hdr[1].text = "Mô tả"
        hdr[2].text = "PIs liên quan"
        for plo in plos:
            _add_table_row(tbl, [plo.id, plo.description, ", ".join(plo.pi_ids)])
        doc.add_paragraph("")

    # ── Teaching Plan ─────────────────────────────────────────────────────────
    weeks = state.get("teaching_weeks", [])
    if weeks:
        _add_heading(doc, "3. Kế Hoạch Giảng Dạy", level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Tuần"
        hdr[1].text = "Chủ đề"
        hdr[2].text = "CLOs"
        hdr[3].text = "LT (tiết)"
        hdr[4].text = "TH (tiết)"
        for w in weeks:
            _add_table_row(
                tbl,
                [
                    str(w.week),
                    w.topic,
                    ", ".join(w.clo_ids),
                    str(w.lecture_hours),
                    str(w.practice_hours),
                ],
            )
        doc.add_paragraph("")

    # ── Assessment ────────────────────────────────────────────────────────────
    assessment = state.get("assessment_plan", [])
    if assessment:
        _add_heading(doc, "4. Kế Hoạch Đánh Giá", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Hình thức"
        hdr[1].text = "Loại"
        hdr[2].text = "Trọng số"
        hdr[3].text = "CLOs"
        for a in assessment:
            _add_table_row(
                tbl,
                [a.name, a.type, f"{a.weight:.0%}", ", ".join(a.clo_ids)],
            )
        doc.add_paragraph("")

    # ── Save ──────────────────────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"DCCT_{course_code}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    logger.info("[WordGenerator] Saved: %s", output_path)
    return output_path
